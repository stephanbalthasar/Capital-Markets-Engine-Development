# app.py
# EUCapML Case Tutor — University of Bayreuth
# - Free LLM via Groq (llama-3.1-8b/70b-versatile): no credits or payments
# - Web retrieval from EUR-Lex, CURIA, ESMA, BaFin, Gesetze-im-Internet
# - Hidden model answer is authoritative; citations [1], [2] map to sources

import hashlib
import json
import math
import numpy as np
import os
import pathlib
import re
import requests
import streamlit as st
import time

from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.document import Document as _Document
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
from typing import List, Dict, Any, Tuple, Union, IO
from urllib.parse import quote_plus, urlparse

BOOKLET = "assets/EUCapML - Course Booklet.docx"

# ---------------- Build fingerprint (to verify latest deployment) ----------------
APP_HASH = hashlib.sha256(pathlib.Path(__file__).read_bytes()).hexdigest()[:10]

# --------------------------- WORD-ONLY PARSER + CITATIONS ---------------------------

# --- Compact citation formatter (no PDF assumptions) ---
def format_booklet_citation(meta: Dict[str, Any]) -> str:
    """
    Produces labels like:
      - "see Course Booklet para. 27"
      - "see Course Booklet Case Study 30"
      - "see Course Booklet" (fallback)
    """
    paras = meta.get("paras") or []
    cases = meta.get("cases") or []
    case_sec = meta.get("case_section")
    case_n = cases[0] if cases else (case_sec if isinstance(case_sec, int) else None)
    if case_n:
        return f"see Course Booklet Case Study {case_n}"
    if paras:
        return f"see Course Booklet para. {paras[0]}"
    return "see Course Booklet"

# --- Word-based parser: extracts numbered paragraphs and case sections ---

PARA_RE_DOTSAFE = re.compile(r"^(\d{1,4})\b(?!\.)")              # 12 but not 1.1
CASE_RE = re.compile(r"^Case\s*Study\s*(\d{1,4})\b", re.I)

def _iter_block_items(parent):
    """
    Yield each paragraph or table within *parent* in document order.
    Works for the main document and for table cells (nested content).
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported container for block iteration")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def _style_is_heading(p: Paragraph) -> bool:
    try:
        name = (p.style.name or "").lower()
        return ("heading" in name) or ("überschrift" in name)  # German UI
    except Exception:
        return False

def _paragraph_is_para_anchor_by_style(p: Paragraph) -> bool:
    """
    Detect 'numbered paragraph' anchors that are rendered via a paragraph style
    (e.g., 'Standard with Para Numbering') rather than a visible digit.
    We match loosely to be resilient against localized style names.
    """
    try:
        name = (p.style.name or "").lower()
        # English: "para", "number"; German Word UIs often show English custom style names too.
        return ("para" in name and "number" in name)
    except Exception:
        return False

def _first_nonempty_run(p: Paragraph):
    for r in p.runs:
        t = (r.text or "").strip()
        if t:
            return r, t
    return None, ""

def _run_is_bold(r) -> bool:
    # direct bold or bold by character style; run.bold may be True/False/None
    try:
        return bool(r.bold) or (getattr(getattr(r, "font", None), "bold", False) is True)
    except Exception:
        return False

def _paragraph_anchor_number(p: Paragraph) -> int | None:
    """
    Anchor if the first non-empty run is BOLD and consists only of digits (1–4).
    Ignore headings and list-like '1.' forms.
    """
    if _style_is_heading(p):
        return None
    r, txt = _first_nonempty_run(p)
    if not txt or not r:
        return None
    # Allow surrounding whitespace but require pure digits; avoid "1." etc.
    m = re.fullmatch(r"\s*(\d{1,4})\s*", txt)
    if m and _run_is_bold(r):
        return int(m.group(1))
    # Fallback: a plain-text line starting with digits not followed by '.' and not a heading
    return None

def _cell_text(cell: _Cell) -> str:
    # Join all paragraphs; preserve spaces
    parts = []
    for para in cell.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    return " ".join(parts).strip()

def _table_row_anchor_number(row) -> int | None:
    """
    Detect a left-cell bold integer (1–4 digits) used as the paragraph marker.
    """
    if not row.cells:
        return None
    left = row.cells[0]
    if not left.paragraphs:
        return None
    p0 = left.paragraphs[0]
    r, txt = _first_nonempty_run(p0)
    if not txt:
        # also allow left cell text like "  12  " without explicit runs
        txt = (p0.text or "").strip()
    m = re.fullmatch(r"\s*(\d{1,4})\s*", txt)
    if m and (r is None or _run_is_bold(r)):  # if there is a run, require bold
        return int(m.group(1))
    return None

def _row_is_empty(row) -> bool:
    return all(not _cell_text(c) for c in row.cells)

def _flush(current: Dict[str, Any], buf: List[str],
           out_chunks: List[str], out_metas: List[Dict[str, Any]]) -> None:
    if not current or not buf:
        return
    text = re.sub(r"\s+", " ", " ".join(buf)).strip()
    if not text:
        return
    meta = {
        "paras": current.get("paras", []),
        "cases": current.get("cases", []),
        "case_section": current.get("case_section"),
    }
    # running index for downstream grouping/sorting
    meta["page_num"] = len(out_metas) + 1
    if meta["cases"]:
        k = meta["cases"][0]
        meta["title"] = f"see Course Booklet Case Study {k}"
        meta["url"]   = f"booklet+docx://case/{k}"
    elif meta["paras"]:
        n = meta["paras"][0]
        meta["title"] = f"see Course Booklet para. {n}"
        meta["url"]   = f"booklet+docx://para/{n}"
    else:
        meta["title"] = "see Course Booklet"
        meta["url"]   = "booklet+docx://booklet"
    out_chunks.append(text)
    out_metas.append(meta)

def parse_booklet_docx(docx_source: Union[str, IO[bytes]]
                       ) -> Tuple[List[str], List[Dict[str, Any]]]:
    """
    Parse the Course Booklet .docx into (chunks, metas) with robust detection:
      • Case Study sections: lines starting with "Case Study <N>"
      • Paragraph anchors: (a) paragraph style-based numbering (e.g., "Standard with Para Numbering"),
                           (b) visible bold integer at line-start,
                           (c) left table-cell integer
      • Continuation: collect subsequent lines until next anchor or next Case Study
      • Headings are not treated as paragraph numbers
    """
    doc = Document(docx_source)

    chunks: List[str] = []
    metas: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {}
    buf: List[str] = []
    current_case: int | None = None

    # Running paragraph id to assign for style-based numbering when no explicit digits are present.
    next_para_id: int = 0

    def start_case(k: int, heading_text: str = ""):
        nonlocal current, buf, current_case
        _flush(current, buf, chunks, metas)
        current = {"cases": [k], "paras": [], "case_section": k}
        buf = [heading_text.strip()] if heading_text.strip() else []
        current_case = k

    def start_para(n: int, first_line: str):
        nonlocal current, buf
        _flush(current, buf, chunks, metas)
        current = {"paras": [n], "cases": [], "case_section": current_case}
        buf = [first_line.strip()] if first_line.strip() else []

    # Iterate paragraphs and tables in document order
    for block in _iter_block_items(doc):

        # ----------------- Plain paragraphs -----------------
        if isinstance(block, Paragraph):
            t = (block.text or "").strip()
            if not t:
                continue

            # Case Study heading outside tables
            m_case = CASE_RE.match(t)
            if m_case:
                start_case(int(m_case.group(1)), t)
                continue

            # (A) Style-based paragraph anchor ("Standard with Para Numbering", etc.)
            if _paragraph_is_para_anchor_by_style(block) and not _style_is_heading(block):
                # Prefer explicit visible integer if present; otherwise assign the next sequential id.
                explicit_n = _paragraph_anchor_number(block)
                if explicit_n is not None:
                    para_id = explicit_n
                    if explicit_n > next_para_id:
                        next_para_id = explicit_n
                else:
                    next_para_id += 1
                    para_id = next_para_id

                start_para(para_id, t)
                continue

            # (B) Visible bold integer at the start (your existing detection)
            n = _paragraph_anchor_number(block)
            if n is not None:
                # If first run is just the number, drop it from the line
                r, _txt = _first_nonempty_run(block)
                rest = t
                if r and re.fullmatch(r"\s*\d{1,4}\s*", r.text or ""):
                    rest = (t[len(r.text):] or "").strip() or t

                if n > next_para_id:
                    next_para_id = n
                start_para(n, rest or t)
                continue

            # Otherwise, continuation of current chunk
            if current:
                buf.append(t)
            continue

        # ----------------- Tables -----------------
        if isinstance(block, Table):
            for row in block.rows:
                if _row_is_empty(row):
                    continue

                left_txt  = _cell_text(row.cells[0]) if len(row.cells) > 0 else ""
                right_txt = _cell_text(row.cells[1]) if len(row.cells) > 1 else ""

                # Case Study heading in a table cell?
                m_case = CASE_RE.match(left_txt) or CASE_RE.match(right_txt)
                if m_case:
                    start_case(int(m_case.group(1)), left_txt or right_txt)
                    continue
                # (A) Style-based para anchor inside any cell of the row
                style_anchor_in_row = False
                if row.cells:
                    for c in row.cells:
                        for p in c.paragraphs:
                            if _paragraph_is_para_anchor_by_style(p) and not _style_is_heading(p):
                                style_anchor_in_row = True
                                break
                        if style_anchor_in_row:
                            break
                if style_anchor_in_row:
                    next_para_id += 1
                    first_line = (right_txt or left_txt).strip()
                    start_para(next_para_id, first_line)
                    continue

                # (B) Numeric anchor in left cell (your existing table detection)
                n = _table_row_anchor_number(row)
                if n is not None:
                    first_line = right_txt.strip()
                    if not first_line:
                        # remove leading digits from left cell if that's all we have
                        first_line = re.sub(r"^\s*\d{1,4}\s*", "", left_txt).strip()
                    if n > next_para_id:
                        next_para_id = n
                    start_para(n, first_line)
                    continue

                # Otherwise treat this row as continuation
                cont = " ".join([x for x in [right_txt, left_txt] if x]).strip()
                if cont and current:
                    buf.append(cont)

    _flush(current, buf, chunks, metas)
    return chunks, metas

# --- Cached loader for parsed anchors (Word) ---
@st.cache_data(show_spinner=False)
def load_booklet_anchors(docx_source: Union[str, IO[bytes]]) -> Tuple[List[Dict[str, Any]], List[str], List[Dict[str, Any]]]:
    """
    Returns:
      records: [{'idx', 'kind', 'id', 'preview'}] where kind ∈ {'para','case','other'}
      chunks:  raw text chunks (same order)
      metas:   per-chunk metadata (same order)
    """
    chunks, metas = parse_booklet_docx(docx_source)

    def first_words(text: str, n=10) -> str:
        toks = (text or "").split()
        return " ".join(toks[:n])

    records: List[Dict[str, Any]] = []
    for idx, (txt, meta) in enumerate(zip(chunks, metas), start=1):
        if meta.get("paras"):
            kind, ident = "para", meta["paras"][0]
        elif meta.get("cases"):
            kind, ident = "case", meta["cases"][0]
        else:
            kind, ident = "other", None
        records.append({
            "idx": idx,
            "kind": kind,
            "id": ident,
            "preview": first_words(txt, 10)
        })
    return records, chunks, metas

# ---------- Public helpers you will call from the app ----------
def _time_budget(seconds: float):
    start = time.monotonic()
    return lambda: time.monotonic() - start < seconds

def normalize_headings(text: str) -> str:
    """Standardize and bold section headings."""
    if not text:
        return text

    # Canonicalize variants
    text = re.sub(r"(?im)^\\s*CLAIMS\\s*:\\s*$", "Student's Core Claims:", text)
    text = re.sub(r"(?im)^\\s*MISTAKES\\s*:\\s*$", "Mistakes:", text)

    # Bold-format canonical headings
    patterns = {
        r"(?im)^\\s*Student's Core Claims:\\s*$": "**Student's Core Claims:**",
        r"(?im)^\\s*Mistakes:\\s*$": "**Mistakes:**",
        r"(?im)^\\s*Missing Aspects:\\s*$": "**Missing Aspects:**",
        r"(?im)^\\s*Suggestions:\\s*$": "**Suggestions:**",
        r"(?im)^\\s*Conclusion:\\s*$": "**Conclusion:**",
    }
    for pat, repl in patterns.items():
        text = re.sub(pat, repl + "\n", text)

    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    return text

# --- Grounding guard helpers (add once) ---
def _anchors_from_model(model_answer_slice: str, cap: int = 20) -> list[str]:
    s = model_answer_slice or ""
    acr = re.findall(r"\b[A-ZÄÖÜ]{2,6}\b", s)                        # MAR, PR, TD, WpHG, ...
    art = re.findall(r"(?i)\b(?:Art\.?|Article)\s*\d+(?:\([^)]+\))*", s)
    par = re.findall(r"§\s*\d+[a-z]?(?:\([^)]+\))*", s)
    # de‑duplicate, keep by original order
    seen, out = set(), []
    for t in acr + art + par:
        t = re.sub(r"\s+", " ", t.strip())
        if t and t not in seen:
            seen.add(t); out.append(t)
        if len(out) >= cap:
            break
    return out

def prune_redundant_improvements(student_answer: str, reply: str, rubric: dict) -> str:
    """
    Removes 'Missing Aspects' bullets that recommend adding content already present
    in the student's answer, based on rubric-derived keyword hits.
    """
    if not reply or not rubric:
        return reply

    # Collect all keywords already detected in the student's answer
    present_keywords = {
        kw.lower().strip()
        for row in rubric.get("per_issue", [])
        for kw in row.get("keywords_hit", [])
        if kw
    }

    # Find the 'Missing Aspects' section
    m = re.search(r"(?is)(Missing Aspects:\s*)(.*?)(?=\n(?:Conclusion|Suggestions|Sources used|$))", reply)
    if not m:
        return reply

    head, body = m.group(1), m.group(2)
    bullets = [ln.strip() for ln in re.split(r"\n\s*•\s*", body.strip()) if ln.strip()]
    kept = []

    for bullet in bullets:
        bullet_lower = bullet.lower()
        # Drop bullet if any present keyword appears in it
        if any(kw in bullet_lower for kw in present_keywords):
            continue
        kept.append(f"• {bullet}")

    new_block = "\n".join(kept) if kept else "—"
    return reply.replace(m.group(0), head + new_block + "\n")

# ---------------- Embeddings ----------------
@st.cache_resource(show_spinner=False)
def load_embedder():
    """
    Try a small sentence-transformer; if unavailable (e.g., install timeouts), fall back to TF-IDF.
    """
    try:
        model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
        return ("sbert", model)
    except Exception:
        return ("tfidf", None)

def embed_texts(texts: List[str], backend):
    kind, model = backend
    if kind == "sbert":
        return model.encode(texts, normalize_embeddings=True)
    vec = TfidfVectorizer(stop_words="english")
    X = vec.fit_transform(texts)
    A = X.toarray()
    norms = np.linalg.norm(A, axis=1, keepdims=True) + 1e-12
    return A / norms

def cos_sim(a, b):
    return float(cosine_similarity(a.reshape(1, -1), b.reshape(1, -1))[0, 0])

DEFAULT_WEIGHTS = {"similarity": 0.4, "coverage": 0.6}

def split_into_chunks(text: str, max_words: int = 180):
    words = text.split()
    chunks, cur = [], []
    for w in words:
        cur.append(w)
        if len(cur) >= max_words:
            chunks.append(" ".join(cur)); cur = []
    if cur: chunks.append(" ".join(cur))
    return chunks

# ---------------- Case & Model Answer (YOUR CONTENT) ----------------
CASE = """
Neon AG is a German stock company (Aktiengesellschaft), the shares of which have been admitted to trading on the regulated market of the Frankfurt stock exchange for a number of years. Gerry is Neon’s CEO (Vorstandsvorsitzender) and holds 25% of Neon’s shares. Gerry wants Neon to develop a new business strategy. For this, Neon would have to buy IP licences for 2.5 billion euros but has no means to afford this. Unicorn plc is a competitor of Neon’s based in the UK and owns licences of the type needed for Neon’s plans. After confidential negotiations, Unicorn, Neon, and Gerry in his personal capacity enter into a “Cooperation Framework Agreement” (“CFA”) which names all three as parties and which has the following terms:
1. Unicorn will transfer the licences to Neon by way of a capital contribution in kind (Sacheinlage). In return, Neon will increase its share capital by 30.1% and issue the new shares to Unicorn. The parties agree that the capital increase should take place within the next 6 months. 
2. Unicorn and Gerry agree that, once the capital increase is complete, they will pre-align major decisions impacting Neon’s business strategy. Where they cannot agree on a specific measure, Gerry agrees to follow Unicorn’s instructions when voting at a shareholder meeting of Neon.
As a result of the capital increase, Gerry will hold approximately 19% in Neon, and Unicorn 23%. Unicorn, Neon and Gerry know that the agreement will come as a surprise to Neon’s shareholders, in particular, because in previous public statements, Gerry had always stressed that he wanted Neon to remain independent. They expect that the new strategy is a “game-changer” for Neon and will change its strategic orientation permanently in a substantial way. 

Questions:
1. Does the conclusion of the CFA trigger capital market disclosure obligations for Neon? What is the timeframe for disclosure? Is there an option for Neon to delay disclosure?
2. Unicorn wants the new shares to be admitted to trading on the regulated market in Frankfurt. Does this require a prospectus under the Prospectus Regulation? What type of information in connection with the CFA would have to be included in such a prospectus?
3. What are the capital market law disclosure obligations that arise for Unicorn once the capital increase and admission to trading are complete and Unicorn acquires the new shares? Can Unicorn participate in Neon’s shareholder meetings if it does not comply with these obligations?

Note:
Your answer will not have to consider the SRD, §§ 111a–111c AktG, or EU capital market law that is not included in your permitted material. You may assume that Gerry and Neon have all corporate authorisations for the conclusion of the CFA and the capital increase.
"""

MODEL_ANSWER = """

1.  Question 1 requires a discussion of whether the conclusion of the CFA triggers an obligation to publish “inside information” pursuant to article 17(1) MAR. 
a)  On the facts of the case (i.e., new shareholder structure of Neon, combined influence of Gerry and Unicorn, substantial change of strategy, etc.), students would have to conclude that the conclusion of the CFA is inside information within the meaning of article 7(1)(a): 
aa) It relates to an issuer (Neon) and has not yet been made public.
bb) Even if the agreement depends on further implementing steps, it creates information of a precise nature within the meaning of article 7(2) MAR in that there is an event that has already occurred – the conclusion of the CFA –, which is sufficient even if one considered it only as an “intermediate step” of a “protracted process”. In addition, subsequent events – the capital increase – can “reasonably be expected to occur” and therefore also qualify as information of a precise nature. A good answer would discuss the “specificity” requirement under article 7(2) MAR and mention that pursuant to the ECJ decision in Lafonta, it is sufficient for the information to be sufficiently specific to constitute a basis on which to assess the effect on the price of the financial instruments, and that the only information excluded by the specificity requirement is information that is “vague or general”. Also, the information is something a reasonable investor would likely use, and therefore likely to have a significant effect on prices within the meaning of article 7(4) MAR.
cc) The information “directly concerns” the issuer in question. As a result, article 17(1) MAR requires Neon to “inform the public as soon as possible”. Students should mention that this allows issuers some time for fact-finding, but otherwise, immediate disclosure is required. Delay is only possible under article 17(4) MAR. However, there is nothing to suggest that Neon has a legitimate interest within the meaning of article 17(4)(a), and at any rate, given previous communication by Neon, a delay would be likely to mislead the public within the meaning of article 17(4)(b). Accordingly, a delay could not be justified under article 17(4) MAR.
Students are not expected to address §§ 33, 38 WpHG. In fact, subscribing to new shares not yet issued (as done in the CFA) does not trigger any disclosure obligations under §§ 38(1), 33(3) WpHG. At any rate, these would only be incumbent on Unicorn, not Neon. 

2.  Question 2 requires an analysis of prospectus requirements under the Prospectus Regulation.
a)  There is no public offer within the meaning of article 2(d) PR that would trigger a prospectus requirement under article 3(1) PR. However, pursuant to article 3(3) PR, admission of securities to trading on a regulated market requires prior publication of a prospectus. Neon shares qualify as securities under article 2(a) PR in conjunction with article 4(1)(44) MiFID II. Students should discuss the fact that there is an exemption for this type of transaction under article 1(5)(a) PR, but that the exemption is limited to a capital increase of 20% or less so does not cover Neon’s case. Accordingly, admission to trading requires publication of a prospectus (under article 21 PR), which in turn makes it necessary to have the prospectus approved under article 20(1) PR). A very complete answer would mention that Neon could benefit from the simplified disclosure regime for secondary issuances under article 14(1)(a) PR.
b)  As regards the content of the prospectus, students are expected to explain that the prospectus would have to include all information in connection with the CFA that is material within the meaning of article 6(1) PR, in particular, as regards the prospects of Neon (article 6(1)(1)(a) PR) and the reasons for the issuance (article 6(1)(1)(c) PR). The prospectus would also have to describe material risks resulting from the CFA and the new strategy (article 16(1) PR). A good answer would mention that the “criterion” for materiality under German case law is whether an investor would “rather than not” use the information for the investment decision.

3.  The question requires candidates to address disclosure obligations under the Transparency Directive and the Takeover Bid Directive and implementing domestic German law. 
a)  As Neon’s shares are listed on a regulated market, Neon is an issuer within the meaning of § 33(4) WpHG, so participations in Neon are subject to disclosure under §§33ff. WpHG. Pursuant to § 33(1) WpHG, Unicorn will have to disclose the acquisition of its stake in Neon. The relevant position to be disclosed includes the 23% stake held by Unicorn directly. In addition, Unicorn will have to take into account Gerry’s 19% stake if the CFA qualifies as “acting in concert” within the meaning of § 34(2) WpHG. In this context, students should differentiate between the two types of acting in concert, namely (i) an agreement to align the exercise of voting rights which qualifies as acting in concert irrespectively of the impact on the issuer’s strategy, and (ii) all other types of alignment which only qualify as acting in concert if it is aimed at modifying substantially the issuer’s strategic orientation. On the facts of the case, both requirements are fulfilled. A good answer should discuss this in the light of the BGH case law, and ideally also consider whether case law on acting in concert under WpÜG can and should be used to assess acting in concert under WpHG. A very complete answer would mention that Unicorn also has to make a statement of intent pursuant to § 43(1) WpHG.
b)  The acquisition of the new shares is also subject to WpÜG requirements pursuant to § 1(1) WpÜG as the shares issued by Neon are securities within the meaning of § 2(2) WpÜG and admitted to trading on a regulated market. Pursuant to § 35(1)(1) WpÜG, Unicorn has to disclose the fact that it acquired “control” in Neon and publish an offer document submit a draft offer to BaFin, §§ 35(2)(1), 14(2)(1) WpÜG. “Control” is defined as the acquisition of 30% or more in an issuer, § 29(2) WpÜG. The 23% stake held by Unicorn directly would not qualify as “control" triggering a mandatory bid requirement. However, § 30(2) WpÜG requires to include in the calculation shares held by other parties with which Unicorn is acting in concert, i.e., Gerry’s 19% stake (students may refer to the discussion of acting in concert under § 34(2) WpHG). The relevant position totals 42% and therefore the disclosure requirements under § 35(1) WpÜG.
c)  Failure to disclose under § 33 WpHG/§ 35 WpÜG will suspend Unicorn’s shareholder rights under § 44 WpHG, § 59 WpÜG. No such sanction exists as regards failure to make a statement of intent under § 43(1) WpHG.
"""

# ---------------- Scoring Rubric ----------------

# ---------- Helpers for robust JSON extraction ----------
def _first_json_block(s: str):
    """Extract the first JSON object/array from a string (handles ```json ... ```)."""
    if not s:
        return None
    m = re.search(r"```json\s*(\{.*?\}|\[.*?\])\s*```", s, flags=re.S | re.I)
    if m:
        return m.group(1)
    m = re.search(r"(\{.*?\}|\[.*?\])", s, flags=re.S)
    return m.group(1) if m else None

def _coerce_issues(parsed) -> list[dict]:
    """Validate & normalize issues into [{name, keywords, importance}] with sane bounds."""
    if parsed is None:
        return []
    issues = parsed.get("issues") if isinstance(parsed, dict) else parsed
    if not isinstance(issues, list):
        return []
    out = []
    for it in issues:
        if not isinstance(it, dict):
            continue
        name = (it.get("name") or "").strip()
        kws  = it.get("keywords") or []
        try:
            importance = int(it.get("importance", 5))
        except Exception:
            importance = 5
        if not name or not isinstance(kws, list):
            continue
        kws = [k.strip() for k in kws if isinstance(k, str) and k.strip()]
        kws = list(dict.fromkeys(kws))[:8]           # dedupe + cap
        importance = max(1, min(10, importance))     # clamp
        if not kws:
            continue
        out.append({"name": name, "keywords": kws, "importance": importance})
    # keep 4–10 issues if possible
    return out[:10]

def _try_parse_json(raw: str):
    """Parse JSON from raw LLM output, being tolerant to fences."""
    if not raw:
        return None
    block = _first_json_block(raw)
    try:
        if block:
            return json.loads(block)
        return json.loads(raw)
    except Exception:
        return None
    
# ---------- Main extractor (no hard-coded topics) ----------
def improved_keyword_extraction(text: str, max_keywords: int = 20) -> list[str]:
    if not text:
        return []

    # Extract law acronyms
    acronyms = re.findall(r"\b(?:MAR|PR|MiFID II|TD|WpHG|WpÜG)\b", text)

    # Extract articles with law names (e.g. "article 7(1) MAR")
    articles = re.findall(
        r"(?i)\b(?:Art\.?|Article)\s*\d+(?:\([^)]+\))*\s*(MAR|PR|MiFID II|TD|WpHG|WpÜG)",
        text
    )
    article_matches = re.findall(
        r"(?i)\b(?:Art\.?|Article)\s*\d+(?:\([^)]+\))*",
        text
    )
    full_articles = []
    for match in article_matches:
        for law in acronyms:
            if law in text:
                full_articles.append(f"{match.strip()} {law}")
                break
    # Extract paragraphs with law names (e.g. "§ 33 WpHG")
    paragraph_matches = re.findall(r"§\s*\d+[a-z]?(?:\([^)]+\))*", text)
    full_paragraphs = []
    for match in paragraph_matches:
        for law in acronyms:
            if law in text and law in ["WpHG", "WpÜG"]:
                full_paragraphs.append(f"{match.strip()} {law}")
                break

    # Extract ECJ cases and named cases
    cases = re.findall(r"C[-–—]?\d+/\d+", text)
    named_cases = re.findall(r"\bLafonta\b|\bGeltl\b|\bHypo Real Estate\b", text, flags=re.I)

    # Combine legal anchors
    legal_anchors = full_articles + full_paragraphs + acronyms + cases + named_cases

    # TF-IDF for generic legal terms
    vec = TfidfVectorizer(ngram_range=(1, 3), max_features=3000, stop_words="english")
    X = vec.fit_transform([text])
    terms = vec.get_feature_names_out()
    tfidf_scores = X.toarray().flatten()

    # Filter out trivial or malformed terms
    blacklist = {
        "requires", "pursuant", "students", "question", "mention", "meaning",
        "agreement", "shares", "neon", "gerry", "company", "framework", "cfa"
    }
    generic_terms = []
    for term, score in zip(terms, tfidf_scores):
        term_clean = term.strip().lower()
        if len(term_clean) < 3 or term_clean in blacklist:
            continue
        if re.search(r"\b(?:requires|pursuant|students|question|mention|meaning)\b", term_clean):
            continue
        generic_terms.append((term, score))

    generic_terms_sorted = sorted(generic_terms, key=lambda x: -x[1])
    top_generic_terms = [term for term, _ in generic_terms_sorted[:max_keywords]]

    # Combine and deduplicate
    all_keywords = legal_anchors + top_generic_terms
    seen = set()
    final_keywords = []
    for kw in all_keywords:
        kw_norm = re.sub(r"\s+", " ", kw.strip())
        if kw_norm.lower() not in seen:
            seen.add(kw_norm.lower())
            final_keywords.append(kw_norm)

    return final_keywords[:max_keywords]
    
def extract_issues_from_model_answer(model_answer: str, llm_api_key: str) -> list[dict]:
    model_answer = (model_answer or "").strip()
    if not model_answer:
        return []

    # Attempt LLM extraction
    sys = "Respond with VALID JSON only: either an array or {\"issues\": [...]}. No prose, no fences."
    user = (
        "Extract the key issues from the MODEL ANSWER.\n"
        "Return JSON ONLY (no code fences). Each item:\n"
        "{ \"name\": \"short issue name\", \"keywords\": [\"3-8 indicative phrases\"], \"importance\": 1-10 }\n\n"
        "MODEL ANSWER:\n" + model_answer
    )
    messages = [
        {"role": "system", "content": sys},
        {"role": "user", "content": user},
    ]

    raw = call_groq(messages, api_key=llm_api_key, model_name=SELECTED_MODEL, temperature=0.0, max_tokens=900)
    parsed = _try_parse_json(raw)
    issues = _coerce_issues(parsed)
            
    # Fallback to improved keyword extraction
    if not issues:
        keywords = improved_keyword_extraction(model_answer, max_keywords=20)
        issues = [{
            "name": "Key Legal Concepts",
            "keywords": keywords,
            "importance": 10
        }]

    return issues

def generate_rubric_from_model_answer(student_answer: str, model_answer: str, backend, llm_api_key: str, weights: dict) -> dict:
    extracted_issues = extract_issues_from_model_answer(model_answer, llm_api_key)
    if not extracted_issues:
        return {
            "similarity_pct": 0.0,
            "coverage_pct": 0.0,
            "final_score": 0.0,
            "per_issue": [],
            "missing": [],
            "bonus": [],
            "substantive_flags": []
        }

    embs = embed_texts([student_answer, model_answer], backend)
    sim = cos_sim(embs[0], embs[1])
    sim_pct = max(0.0, min(100.0, 100.0 * (sim + 1) / 2))

    per_issue, tot, got = [], 0, 0
    bonus = []

    for issue in extracted_issues:
        pts = issue.get("importance", 5) * 2
        tot += pts
        sc, hits = coverage_score(student_answer, {"keywords": issue["keywords"], "points": pts})
        got += sc
        per_issue.append({
            "issue": issue["name"],
            "max_points": pts,
            "score": sc,
            "keywords_hit": hits,
            "keywords_total": issue["keywords"],
        })
        
    cov_pct = 100.0 * got / max(1, tot)
    final = (weights["similarity"] * sim_pct + weights["coverage"] * cov_pct) / (weights["similarity"] + weights["coverage"])

    # Only required issues can be marked as "missing"
    missing = []
    for row in per_issue:
        missed = [kw for kw in row["keywords_total"] if not keyword_present(student_answer, kw)]
        if missed:
            missing.append({"issue": row["issue"], "missed_keywords": missed})

    substantive_flags = detect_substantive_flags(student_answer)

    return {
        "similarity_pct": round(sim_pct, 1),
        "coverage_pct": round(cov_pct, 1),
        "final_score": round(final, 1),
        "per_issue": per_issue,
        "missing": missing,
        "bonus": bonus,
        "substantive_flags": substantive_flags,
    }

def filter_model_answer_and_rubric(selected_question: str, model_answer: str, api_key: str) -> tuple[str, list[dict]]:
    # Find section starts anchored at line-beginnings
    m1 = re.search(r"(?m)^\s*1\.\s", model_answer)
    m2 = re.search(r"(?m)^\s*2\.\s", model_answer)
    m3 = re.search(r"(?m)^\s*3\.\s", model_answer)

    start1 = m1.start() if m1 else 0
    start2 = m2.start() if m2 else None
    start3 = m3.start() if m3 else None
    end = len(model_answer)

    if selected_question == "Question 1":
        lo, hi = start1, (start2 if start2 is not None else end)
    elif selected_question == "Question 2":
        lo, hi = (start2 if start2 is not None else 0), (start3 if start3 is not None else end)
    elif selected_question == "Question 3":
        lo, hi = (start3 if start3 is not None else 0), end
    else:
        lo, hi = 0, end

    model_answer_filtered = model_answer[lo:hi].strip()
    extracted_issues = extract_issues_from_model_answer(model_answer_filtered, api_key)
    return model_answer_filtered, extracted_issues
    
# ---------------- Robust keyword & citation checks ----------------
def canonicalize(s: str, strip_paren_numbers: bool = False) -> str:
    s = s.lower()
    s = s.replace("art.", "art").replace("article", "art").replace("–", "-")
    s = s.replace("wpüg", "wpüg")
    s = re.sub(r"\s+", "", s)
    if strip_paren_numbers:
        s = re.sub(r"\(\d+[a-z]?\)", "", s)
    s = re.sub(r"[^a-z0-9§]", "", s)
    return s

def keyword_present(answer: str, kw: str) -> bool:
    """
    Detects presence of compound legal references like 'article 17(4)(a) MAR' or '§ 33 WpHG'
    even if the student mentions the article/paragraph and the law name separately.
    """
    
    def canonicalize(s: str) -> str:
        s = s.lower()
        s = s.replace("art.", "art").replace("article", "art").replace("–", "-")
        s = s.replace("wpüg", "wpüg")
        s = re.sub(r"\s+", "", s)
        s = re.sub(r"[^\w§]", "", s)
        return s

    ans_can = canonicalize(answer)
    kw_can = canonicalize(kw)

    # Direct match
    if kw_can in ans_can:
        return True

    # Split compound keywords like 'article 17(4)(a) MAR'
    parts = kw.strip().split()
    if len(parts) >= 2:
        main_part = " ".join(parts[:-1])
        law_part = parts[-1]
        main_can = canonicalize(main_part)
        law_can = canonicalize(law_part)
        return main_can in ans_can and law_can in ans_can

    # Fallback: check if keyword appears loosely
    return canonicalize(kw) in ans_can

def coverage_score(answer: str, issue: Dict) -> Tuple[int, List[str]]:
    hits = [kw for kw in issue["keywords"] if keyword_present(answer, kw)]
    score = int(round(issue["points"] * (len(hits) / max(1, len(issue["keywords"])))))
    return score, hits

def detect_substantive_flags(answer: str) -> List[str]:
    flags = []
    low = answer.lower()
    if "always delay" in low or re.search(r"\b(can|may)\s+always\s+delay\b", low):
        flags.append("Delay under Art 17(4) MAR is conditional: (a) legitimate interest, (b) not misleading, (c) confidentiality ensured.")
    return flags

def _find_section(text: str, title_regex: str):
    """
    Return (head, body, tail, span) for the section whose title matches title_regex.
    If not found, returns (None, None, None, None).
    """
    m = re.search(
        rf"({title_regex}\s*)(.*?)(\n(?:Student's Core Claims:|Mistakes:|Missing Aspects:|Suggestions:|Conclusion|📚|Sources used|$))",
        text,
        flags=re.S | re.I,
    )
    if not m:
        return None, None, None, None
    return m.group(1), m.group(2), m.group(3), m.span(0)

def _neutralise_error_tone(line: str) -> str:
    """
    Turn blamey phrasing into 'suggestion' tone.
    """
    s = line
    s = re.sub(r"\b[Tt]he student incorrectly (states|assumes|concludes)\b", "Consider also", s)
    s = re.sub(r"\b[Tt]his is incorrect because\b", "Rationale:", s)
    s = s.replace("is incorrect", "may be incomplete")
    return s

def tidy_empty_sections(reply: str) -> str:
    """
    Remove headings that ended up empty after normalisation.
    """
    if not reply:
        return reply
    reply = normalize_headings(reply)
    return reply

def format_feedback_and_filter_missing(reply: str, student_answer: str, model_answer_slice: str, rubric: dict) -> str:
    """
    Reformats feedback into five clear sections:
    - Student's Core Claims
    - Mistakes
    - Missing Aspects
    - Suggestions
    - Conclusion
    Each section is formatted with bullet points and explanations where needed.
    """
    
    if not reply:
        return reply

        reply = normalize_headings(reply)
    
    # Remove hallucinated 'Missing Aspects' (already present in student answer)
    present = set()
    for row in (rubric or {}).get("per_issue", []):
        present.update({kw.lower() for kw in row.get("keywords_hit", [])})

    def _find_section(text, title_regex):
        m = re.search(rf"({title_regex}\\s*)(.*?)(\\n(?:\\*\\*Student's Core Claims:\\*\\*|\\*\\*Mistakes:\\*\\*|\\*\\*Suggestions:\\*\\*|\\*\\*Conclusion:\\*\\*|$))", text, flags=re.S | re.I)
        return m.groups() if m else (None, None, None)

    present_keywords = {kw.lower() for row in rubric.get("per_issue", []) for kw in row.get("keywords_hit", [])}
    
    head, body, tail = _find_section(reply, r"\\*\\*Missing Aspects:\\*\\*")
    if head:
        bullets = [f"• {ln.strip()}" for ln in body.strip().splitlines() if ln.strip()]
        def fuzzy_keyword_match(bullet: str, present_keywords: set) -> bool:
            bullet_low = bullet.lower()
            for kw in present_keywords:
                kw_tokens = kw.lower().split()
                if all(tok in bullet_low for tok in kw_tokens):
                    return True
            return False
        kept = [b for b in bullets if not any(p in b.lower() for p in present)]
        reply = reply.replace(head + body + tail, head + ("\n".join(kept) + "\n" if kept else "—\n") + tail)

    # Collapse excessive blank lines
    reply = re.sub(r"\n{3,}", "\n\n", reply).strip()

    return reply

# MODEL-CONSISTENCY GUARDRAIL (general, no question-specific logic)
# =======================

def _json_only(messages, api_key,  model_name=None, max_tokens=700):
    model_name = model_name or SELECTED_MODEL
    raw = call_groq(messages, api_key=api_key, model_name=SELECTED_MODEL, temperature=0.0, max_tokens=max_tokens)
    return _try_parse_json(raw)

def check_reply_vs_model_for_contradictions(model_answer: str, reply: str, api_key: str,  model_name=None) -> dict:
    """
    Structured 'consistency critic' that flags contradictions between ASSISTANT_REPLY and MODEL_ANSWER.
    Returns: {"consistent": bool, "contradictions": [{"reply_span","model_basis","why","fix"}]}
    """
    if not api_key or not reply or not model_answer:
        return {"consistent": True, "contradictions": []}

    ma = truncate_block(model_answer, 3200)
    rp = truncate_block(reply, 1800)

    sys = (
        "You are a strict checker. OUTPUT VALID JSON ONLY (no prose, no fences).\n"
        "Task: Compare the ASSISTANT_REPLY to the AUTHORITATIVE_MODEL_ANSWER.\n"
        "Identify statements in ASSISTANT_REPLY that contradict or materially diverge from the MODEL_ANSWER.\n"
        "Focus on conclusions, rules, tests, thresholds, outcomes. Ignore style. If in doubt, prefer the MODEL_ANSWER."
    )
    user = (
        "{"
        "\"spec\":\"Return JSON: {\\\"consistent\\\":true|false, \\\"contradictions\\\":[{\\\"reply_span\\\":<=200c, \\\"model_basis\\\":<=200c, \\\"why\\\":<=120c, \\\"fix\\\":<=160c}]}\""
        "}\n\n"
        "AUTHORITATIVE_MODEL_ANSWER:\n\"\"\"\n" + ma + "\n\"\"\"\n\n"
        "ASSISTANT_REPLY:\n\"\"\"\n" + rp + "\n\"\"\"\n"
    )
    parsed = _json_only(
        [{"role": "system", "content": sys}, {"role": "user", "content": user}],
        api_key, model_name=model_name, max_tokens=700
    )
    if not parsed or not isinstance(parsed, dict):
        return {"consistent": True, "contradictions": []}

    contrad = parsed.get("contradictions") or []
    clean = []
    for c in contrad:
        if not isinstance(c, dict):
            continue
        rs = (c.get("reply_span") or "")[:220]
        mb = (c.get("model_basis") or "")[:220]
        wy = (c.get("why") or "")[:140]
        fx = (c.get("fix") or "")[:180]
        if rs and mb:
            clean.append({"reply_span": rs, "model_basis": mb, "why": wy, "fix": fx})

    return {"consistent": not bool(clean), "contradictions": clean}

def rewrite_reply_to_match_model(model_answer: str, reply: str, contradictions: list, api_key: str,  model_name=None) -> str:
    """
    Rewrites the reply to align with the MODEL_ANSWER.
    Preserves structure, ≤400 words, keeps existing [n] citations but does NOT invent new numbers.
    """
    if not api_key or not reply or not model_answer:
        return reply

    ma = truncate_block(model_answer, 3200)
    rp = truncate_block(reply, 1800)

    report = "\n".join(
        f"{i}. reply: {c.get('reply_span','')}\n   model: {c.get('model_basis','')}\n   fix: {c.get('fix','')}"
        for i, c in enumerate(contradictions[:8], 1)
    )

    sys = (
        "You are a careful editor. Rewrite ASSISTANT_REPLY so it does NOT contradict the AUTHORITATIVE_MODEL_ANSWER.\n"
        "Keep ≤400 words, preserve structure/voice, and KEEP existing numeric bracket citations [n].\n"
        "Do NOT invent new numbers; you may delete/relocate an inappropriate [n].\n"
        "If the student is wrong per the model, state the correct conclusion first and explain briefly."
    )
    user = (
        "AUTHORITATIVE_MODEL_ANSWER:\n\"\"\"\n" + ma + "\n\"\"\"\n\n"
        "ASSISTANT_REPLY (to correct):\n\"\"\"\n" + rp + "\n\"\"\"\n\n"
        "INCONSISTENCIES:\n" + report + "\n\n"
        "OUTPUT ONLY the corrected reply text (no JSON, no preface)."
    )

    fixed = call_groq(
        [{"role": "system", "content": sys}, {"role": "user", "content": user}],
        api_key=api_key, model_name=model_name, temperature=0.1, max_tokens=900
    )
    return fixed or reply

def enforce_model_consistency(reply: str, model_answer_filtered: str, api_key: str,  model_name=None) -> str:
    """
    Detect → correct → (optionally) verify.
    If LLM unavailable or nothing to fix, returns original reply.
    """
    if not reply:
        return reply

    check = check_reply_vs_model_for_contradictions(model_answer_filtered, reply, api_key, model_name)
    if check.get("consistent", True):
        return reply

    corrected = rewrite_reply_to_match_model(
        model_answer_filtered, reply, check.get("contradictions", []),
        api_key, model_name
    ) or reply

    # Best-effort recheck; keep corrected either way
    recheck = check_reply_vs_model_for_contradictions(model_answer_filtered, corrected, api_key, model_name)
    return corrected if recheck.get("consistent", True) else corrected
    
# ---------------- Web Retrieval (RAG) ----------------
ALLOWED_DOMAINS = {
    "eur-lex.europa.eu",        # EU law (MAR, PR, MiFID II, TD)
    "curia.europa.eu",          # CJEU (Lafonta C‑628/13 etc.)
    "www.esma.europa.eu",       # ESMA guidelines/news
    "www.bafin.de",             # BaFin
    "www.gesetze-im-internet.de", "gesetze-im-internet.de",  # WpHG, WpÜG
    "www.bundesgerichtshof.de", # BGH
}

SEED_URLS = [
    "https://eur-lex.europa.eu/eli/reg/2014/596/oj",   # MAR
    "https://eur-lex.europa.eu/eli/reg/2017/1129/oj",  # Prospectus Regulation
    "https://eur-lex.europa.eu/eli/dir/2014/65/oj",    # MiFID II
    "https://eur-lex.europa.eu/eli/dir/2004/109/oj",   # Transparency Directive
    "https://curia.europa.eu/juris/liste.jsf?num=C-628/13",  # Lafonta
    "https://www.gesetze-im-internet.de/wphg/",
    "https://www.gesetze-im-internet.de/wpu_g/",
    "https://www.esma.europa.eu/press-news/esma-news/esma-finalises-guidelines-delayed-disclosure-inside-information-under-mar",
]

UA = {"User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"}

@st.cache_data(ttl=3600, show_spinner=False)
def duckduckgo_search(query: str, max_results: int = 6) -> List[Dict]:
    url = f"https://html.duckduckgo.com/html/?q={quote_plus(query)}"
    try:
        r = requests.get(url, headers=UA, timeout=6)
        r.raise_for_status()
    except Exception:
        return []
    soup = BeautifulSoup(r.text, "lxml")
    out = []
    for a in soup.select("a.result__a"):
        href = a.get("href")
        title = a.get_text(" ", strip=True)
        if not href:
            continue
        domain = urlparse(href).netloc.lower()
        if any(domain.endswith(d) for d in ALLOWED_DOMAINS):
            out.append({"title": title, "url": href})
        if len(out) >= max_results:
            break
    return out

@st.cache_data(ttl=3600, show_spinner=False)
def fetch_url(url: str) -> Dict:
    try:
        r = requests.get(url, headers=UA, timeout=8)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "lxml")
        for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
            tag.decompose()
        title = (soup.title.get_text(strip=True) if soup.title else url)
        text = " ".join(soup.stripped_strings)
        text = re.sub(r"\s+", " ", text)
        if len(text) > 120000:
            text = text[:120000]
        return {"url": url, "title": title, "text": text}
    except Exception:
        return {"url": url, "title": url, "text": ""}

def build_queries(student_answer: str, 
    extracted_keywords: List[str], 
    extra_user_q: str = "", 
    max_keywords: int = 6, 
    max_domains: int = 3
    ) -> List[str]:
    """
    Dynamically builds search queries for legal sources based on extracted keywords and user input.
    Targets EUR-Lex, CURIA, BaFin, ESMA, Gesetze-im-Internet.

    Args:
        student_answer (str): The student's written answer.
        extracted_keywords (List[str]): Keywords extracted from the model answer or rubric.
        extra_user_q (str): Optional follow-up question from the user.

    Returns:
        List[str]: A list of search queries suitable for DuckDuckGo.
    """
    base_queries = []
    legal_domains = [
        "site:eur-lex.europa.eu",
        "site:curia.europa.eu",
        "site:esma.europa.eu",
        "site:bafin.de",
        "site:gesetze-im-internet.de"
    ]
    # Clean and normalize keywords
    keywords = [kw.strip() for kw in extracted_keywords if kw and len(kw.strip()) >= 3]
    keywords = list(dict.fromkeys(keywords))[:20]  # deduplicate and cap

    # Build queries from keywords
    for kw in keywords:
        for domain in legal_domains:
            base_queries.append(f"{kw} {domain}")

    # Add student answer context if available
    if student_answer:
        context_snippet = student_answer[:300].strip().replace("\n", " ")
        base_queries.append(f"({context_snippet}) {' OR '.join(legal_domains)}")

    # Add extra user question if provided
    if extra_user_q:
        base_queries.append(f"{extra_user_q.strip()} {' OR '.join(legal_domains)}")

    return base_queries

def collect_corpus(student_answer: str,
                   extracted_keywords: List[str],
                   extra_user_q: str,
                   max_fetch: int = 20,
                   search_budget_s: float = 15.0,   # total time for searches
                   fetch_budget_s: float = 12.0,    # total time for page fetches
                   max_workers: int = 8             # concurrency
                   ) -> List[Dict]:

    # Seed URLs first (zero-cost)
    results = [{"title": "", "url": u} for u in SEED_URLS]

    # Build fewer queries
    queries = build_queries(student_answer, extracted_keywords, extra_user_q)

    # ---- Concurrent search with wall-clock budget ----
    within_budget = _time_budget(search_budget_s)
    search_hits = []
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        futs = {ex.submit(duckduckgo_search, q, 5): q for q in queries}
        for fut in as_completed(futs, timeout=search_budget_s + 2):
            if not within_budget():
                break
            try:
                search_hits.extend(fut.result() or [])
            except Exception:
                pass
            if len(search_hits) >= 40:   # soft cap
                break

    results.extend(search_hits)

    # Clean + keep allowed domains
    seen, cleaned = set(), []
    for r in results:
        url = r["url"]
        if url in seen:
            continue
        seen.add(url)
        domain = urlparse(url).netloc.lower()
        if any(domain.endswith(d) for d in ALLOWED_DOMAINS):
            cleaned.append(r)

    # ---- Concurrent page fetch with wall-clock budget ----
    fetched, within_budget = [], _time_budget(fetch_budget_s)
    to_fetch = cleaned[:max_fetch]
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        futs = {ex.submit(fetch_url, r["url"]): r for r in to_fetch}
        for fut in as_completed(futs, timeout=fetch_budget_s + 2):
            if not within_budget():
                break
            try:
                pg = fut.result()
                if pg.get("text"):
                    # carry over title if fetch_url didn't set one
                    r = futs[fut]
                    if not pg.get("title"):
                        pg["title"] = r.get("title") or r["url"]
                    fetched.append(pg)
            except Exception:
                pass
            if len(fetched) >= max_fetch:
                break

    return fetched

# ---- Booklet relevance terms per question ----
def booklet_chunk_relevant(text: str, extracted_keywords: list[str], user_query: str = "") -> bool:
    q_terms = [w.lower() for w in re.findall(r"[A-Za-zÄÖÜäöüß0-9\-]{3,}", user_query or "")]
    keys = [k.lower() for k in (extracted_keywords or [])]
    tgt = text.lower()
    return any(k in tgt for k in (keys + q_terms))

def retrieve_snippets_with_booklet(student_answer, model_answer_filtered, pages, backend,
                                  extracted_keywords, user_query: str = "",
                                  top_k_pages=8, chunk_words=170):
    booklet_chunks, booklet_metas = [], []
    try:
        booklet_chunks, booklet_metas = parse_booklet_docx(docx_source)         
    except Exception as e:
        st.warning(f"Could not load course booklet: {e}")
    try:
        _model_anchors = _anchors_from_model(model_answer_filtered)
    except Exception as _e:
        _model_anchors = []

    if _model_anchors:
        alow = [a.lower() for a in _model_anchors]
        mc2, mm2 = [], []
        for ch, meta in zip(booklet_chunks, booklet_metas):
            txt = (ch or "").lower()
            if any(a in txt for a in alow):
                mc2.append(ch)
                mm2.append(meta)
        if mc2:  # shrink only if something kept
            booklet_chunks, booklet_metas = mc2, mm2

    # (keep the rest unchanged)

    # ✅ Filter booklet chunks using keywords + the user's query AND case numbers, if any
    selected_q = st.session_state.get("selected_question", "Question 1")
    uq_cases = detect_case_numbers(user_query or "")
    filtered_chunks, filtered_metas = [], []
    for ch, m in zip(booklet_chunks, booklet_metas):
        has_kw = booklet_chunk_relevant(ch, extracted_keywords, user_query)
        case_match = bool(uq_cases and set(uq_cases).intersection(set(m.get("cases") or [])))
        if has_kw or case_match:
            filtered_chunks.append(ch)
            filtered_metas.append(m)
    if filtered_chunks:
        booklet_chunks, booklet_metas = filtered_chunks, filtered_metas
    
    # ---- Prepare booklet meta tuples with a unique key per *page* so we can group snippets by page
    booklet_meta = []
    for m in booklet_metas:
        page_key = -(m["page_num"])  
        citation = format_booklet_citation(m)  # pre-format a nice line
        # We store citation in 'title' so we can reuse downstream without new structures
        booklet_meta.append((page_key, "booklet://course-booklet", citation))

    # ---- Prepare web chunks (unchanged)
    # ---- Prepare web chunks (fixed: keep meta 1:1 with chunks)
    web_chunks, web_meta = [], []
    selected_q = st.session_state.get("selected_question", "Question 1")
    for i, p in enumerate(pages):
        text = p.get("text", "")
        if not text:
            continue
    # Optional relevance filter (keep if you added BOOKLET_KEY_TERMS):
        if 'web_page_relevant' in globals() and not web_page_relevant(text, extracted_keywords):
            continue

        chunks_i = split_into_chunks(text, max_words=chunk_words)
        for ch in chunks_i:
            web_chunks.append(ch)
            web_meta.append((i + 1, p["url"], p["title"]))  # append meta PER CHUNK

    # ---- Build combined corpus
    all_chunks = booklet_chunks + web_chunks
    all_meta   = booklet_meta   + web_meta
    # Defensive: keep chunks and meta in lockstep
    if len(all_chunks) != len(all_meta):
        m = min(len(all_chunks), len(all_meta))
        all_chunks = all_chunks[:m]
        all_meta   = all_meta[:m]

    # Query vector built from student + model slice
    query = "\n\n".join([s for s in [user_query, student_answer, model_answer_filtered] if s])
    embs = embed_texts([query] + all_chunks, backend)
    qv, cvs = embs[0], embs[1:]
    sims = [cos_sim(qv, v) for v in cvs]
    idx = np.argsort(sims)[::-1]

    # ✅ Similarity floor to keep only reasonably relevant snippets
    MIN_SIM = 0.22  # tune if needed

    # ---- Select top snippets grouped by (booklet page) or (web page index)
    per_page = {}
    for j in idx:
        if sims[j] < MIN_SIM:
            break
        pi, url, title = all_meta[j]
        snip = all_chunks[j]
        arr = per_page.setdefault(pi, {"url": url, "title": title, "snippets": []})
        if len(arr["snippets"]) < 3:
            arr["snippets"].append(snip)
        if len(per_page) >= top_k_pages:
            break

    # Order by key and build source lines. For booklet items we already have 'title' as a full citation line.
    top_pages = [per_page[k] for k in sorted(per_page.keys())][:top_k_pages]

    source_lines = []
    for i, tp in enumerate(top_pages):
        if tp["url"].startswith("booklet://"):
            # already a fully formatted citation like: "Course Booklet — p. ii (PDF p. 4), para. 115"
            source_lines.append(f"[{i+1}] {tp['title']}")
        else:
            source_lines.append(f"[{i+1}] {tp['title']} — {tp['url']}")

    return top_pages, source_lines

# ---------------- LLM via Groq (free) ----------------
def call_groq(messages: List[Dict], api_key: str, model_name=None,
              temperature: float = 0.2, max_tokens: int = 700) -> str:
    """
    Groq OpenAI-compatible chat endpoint. Models like llama-3.1-8b-instant / 70b-instant are free.
    """
    if not api_key:
        st.error("No GROQ_API_KEY found (add it to Streamlit Secrets).")
        return None
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {
        "model": model_name,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    try:
        r = requests.post(url, headers=headers, json=data, timeout=90)
        if r.status_code != 200:
            # Show exact error so it's easy to fix
            try: body = r.json()
            except Exception: body = r.text
            st.error(f"Groq error {r.status_code}: {body}")
            return None
        return r.json()["choices"][0]["message"]["content"]
    except requests.exceptions.Timeout:
        st.error("Groq request timed out (60s). Try again or reduce max_tokens.")
        return None
    except requests.exceptions.RequestException as e:
        st.error(f"Groq request failed: {e}")
        return None

def system_guardrails():
    return (
        "You are a careful EU/German capital markets law tutor.\n"
        "HARD RULES:\n"
        "• Base feedback on the AUTHORITATIVE MODEL_ANSWER; do not contradict it.\n"
        "• If student reasoning conflicts, state the correct conclusion first and explain briefly.\n"
        "• Do not disclose or refer to any internal reference material.\n"
        "\n"
        "CITATIONS:\n"
        "• Cite ONLY using numeric brackets [n] that refer to the provided SOURCES.\n"
        "• Only use [n] if the corresponding EXCERPT [n] actually supports the claim.\n"
        "• Do NOT fabricate Course Booklet references.\n"
        "\n"
        "FEEDBACK PRINCIPLES:\n"
        "• If the student substantially aligns with the MODEL_ANSWER, mark claims as Correct and put extras under Suggestions.\n"
        "• If central concepts are missing, explain why they matter; correct mis-citations succinctly.\n"
        "• Summarise—do not copy long passages.\n"
        "\n"
        "OUTPUT:\n"
        "• ≤400 words, concise, didactic, actionable; same language as the student.\n"
        "• Headings must be exactly:\n"
        "**Student's Core Claims:**\n"
        "**Mistakes:**\n"
        "**Missing Aspects:**\n"
        "**Suggestions:**\n"
        "**Conclusion:**\n"
    )

def _flatten_hits_misses_from_rubric(rubric: dict) -> tuple[list[str], list[str]]:
    """
    From the computed rubric, extract:
    - present_keywords: every keyword we *deterministically* detected in the student's answer
    - missing_keywords: every keyword we *did not* detect (flattened from rubric['missing'])
    Both lists are lowercased and de-duplicated.
    """
    present = []
    for row in (rubric or {}).get("per_issue", []):
        present.extend(row.get("keywords_hit", []))
    missing = []
    for m in (rubric or {}).get("missing", []):
        missing.extend(m.get("missed_keywords", []))

    # Normalise / dedupe
    norm = lambda s: re.sub(r"\s+", " ", (s or "").strip()).lower()
    present = list(dict.fromkeys(norm(k) for k in present if k))
    missing = list(dict.fromkeys(norm(k) for k in missing if k))

    # Keep lists reasonably short for the prompt
    return present[:30], missing[:30]

def build_feedback_prompt(student_answer, rubric, model_answer, sources_block, excerpts_block):
    has_sources = sources_block and "(no web sources available)" not in sources_block
    citation_rule = (
        "Use numeric citations [n] only from SOURCES, and only when the EXCERPT [n] supports the claim."
        if has_sources else
        "No numeric citations are available for this reply; write without [n] citations."
    )

    issue_names = [row["issue"] for row in rubric.get("per_issue", [])]
    present = [kw for row in rubric.get("per_issue", []) for kw in row.get("keywords_hit", [])]
    missing = [kw for m in rubric.get("missing", []) for kw in m.get("missed_keywords", [])]
    present_block = "• " + "\n• ".join(present) if present else "—"
    missing_block = "• " + "\n• ".join(missing) if missing else "—"

    return f"""
GRADE THE STUDENT'S ANSWER USING THE RUBRIC AND THE WEB/BOOKLET SOURCES.

STUDENT ANSWER:
\"\"\"{student_answer}\"\"\"

RUBRIC SUMMARY:
- Similarity: {rubric.get('similarity_pct', 0)}%
- Coverage: {rubric.get('coverage_pct', 0)}%
- Overall score: {rubric.get('final_score', 0)}%
- Issues to cover: {", ".join(issue_names) if issue_names else "—"}

MODEL ANSWER (authoritative):
\"\"\"{model_answer}\"\"\"

SOURCES (numbered):
{sources_block}

EXCERPTS (quote sparingly):
{excerpts_block}

AUTO-DETECTED EVIDENCE:
- PRESENT (do NOT mark these as Mistakes/Missing): 
{present_block}
- POTENTIALLY MISSING (mark only if truly absent and material):
{missing_block}

OUTPUT FORMAT (exact headings and bulleting):
**Student's Core Claims:**
• <claim> — [Correct|Incorrect|Not supported]
**Mistakes:**
• <incorrect claim> — short why [n]
**Missing Aspects:**
• <missing concept> — why it matters [n]
**Suggestions:**
• <optional improvements> [n]
**Conclusion:**
<one sentence>

RULES:
- {citation_rule}
- Always insert a line break after each heading and before each bullet.
- Be concise and actionable; ≤400 words total.
""".strip()

def lock_out_false_mistakes(reply: str, rubric: dict) -> str:
    """
    Removes any 'Mistakes' bullet that mentions a keyword already present in the student's answer.
    """
    if not reply or not rubric:
        return reply

    present = {kw.lower() for row in rubric.get("per_issue", []) for kw in row.get("keywords_hit", [])}
    m = re.search(r"(?is)(\*\*Mistakes:\*\*\s*)(.*?)(?=\n\*\*|$)", reply)
    if not m:
        return reply

    head, body = m.group(1), m.group(2)
    bullets = [b.strip() for b in re.split(r"\s*•\s*", body) if b.strip()]
    kept = []

    for b in bullets:
        low = b.lower()
        if any(p in low for p in present):
            continue  # Drop hallucinated mistake
        kept.append(f"• {b}")

    new_block = head + ("\n".join(kept) if kept else "—") + "\n"
    return reply.replace(m.group(0), new_block)

def lock_out_false_missing(reply: str, rubric: dict) -> str:
    """
    Removes any 'Missing Aspects' bullet whose text contains a keyword we already
    detected in the student's answer (rubric['per_issue'][...]['keywords_hit']).
    """
    if not reply:
        return reply

    try:
        present, _ = _flatten_hits_misses_from_rubric(rubric)
        present_set = {p.lower() for p in present}

        # Find the 'Missing Aspects:' section using the same logic as _find_section
        head, body, tail, span = _find_section(reply, r"Missing Aspects:")
        if not head:
            return reply

        # Turn the section body into bullets, filter out any bullet that mentions a present keyword
        bullets = [f"• {ln.strip()}" for ln in body.strip().splitlines() if ln.strip()]
        kept = []
        for b in bullets:
            low = re.sub(r"\s+", " ", b).lower()
            if any(k in low for k in present_set):
                # Drop bullet: it hallucinates "missing" for something we already saw
                continue
            kept.append(b)

        new_block = "—\n" if not kept else "\n".join(kept) + "\n"
        return reply.replace(head + body + tail, head + new_block + tail)
    except Exception:
        return reply

def build_chat_messages(chat_history: List[Dict], model_answer: str, sources_block: str, excerpts_block: str) -> List[Dict]:
    msgs = [{"role": "system", "content": system_guardrails()}]

    # --- Step 3: Hard rule to keep replies aligned with the MODEL ANSWER ---
    msgs.append({"role": "system", "content":
        "HARD RULE: Do not contradict the MODEL ANSWER. "
        "If student reasoning conflicts, state the correct conclusion per the MODEL ANSWER and explain briefly."
    })

    for m in chat_history[-8:]:
        if m.get("role") in ("user", "assistant"):
            msgs.append(m)

    # Pin authoritative context and sources
    msgs.append({"role": "system", "content": "MODEL ANSWER (authoritative):\n" + model_answer})
    msgs.append({"role": "system", "content": "SOURCES:\n" + sources_block})
    msgs.append({"role": "system", "content": "RELEVANT EXCERPTS (quote sparingly):\n" + excerpts_block})
    return msgs

# ------------------------------ Chat and Feebdack Helpers ----------
def web_page_relevant(text: str, extracted_keywords: list[str]) -> bool:
    return any(kw.lower() in text.lower() for kw in extracted_keywords)

# ---- Output completeness helpers ----
def is_incomplete_text(text: str) -> bool:
    """Heuristic: returns True if the text likely ends mid-sentence."""
    if not text or not text.strip():
        return True
    tail = text.strip()[-1]
    return tail not in ".!?…’”\")»]"

def truncate_block(s: str, max_chars: int = 3600) -> str:
    """Trim very long prompt sections to reduce truncation risk."""
    s = s or ""
    return s if len(s) <= max_chars else (s[:max_chars] + " …")

def generate_with_continuation(messages, api_key,  model_name=None, temperature=0.2, first_tokens=1500, continue_tokens=500):
    """
    Calls the LLM, and if output ends mid-sentence, asks it to continue once.
    """
    reply = call_groq(messages, api_key, model_name=model_name, temperature=temperature, max_tokens=first_tokens)
    if reply and is_incomplete_text(reply):
        # Ask for a short continuation to finish the sentence + a 1‑sentence conclusion
        cont_msgs = messages + [{
            "role": "user",
            "content": "Continue exactly where you left off. Finish the previous sentence and add a single-sentence conclusion. Do not repeat earlier text."
        }]
        more = call_groq(cont_msgs, api_key, model_name=model_name, temperature=min(temperature, 0.3), max_tokens=continue_tokens)
        if more:
            reply = (reply.rstrip() + "\n" + more.strip())
    return reply

# --- Citation post-processing & filtering ---
def parse_cited_indices(text: str) -> list[int]:
    try:
        return sorted({int(x) for x in re.findall(r"\[(\d+)\]", text or "")})
    except Exception:
        return []

def filter_sources_by_indices(source_lines: list[str], used: list[int]) -> list[str]:
    """Return only those lines whose [n] was actually cited; preserve numbering."""
    if not used:
        return []
    out = []
    for n in used:
        if 1 <= n <= len(source_lines):
            out.append(source_lines[n - 1])
    return out

# Paragraph markers may appear as "para. 115", "paragraph 115", "Rn. 115", "[115]", "¶ 115"
_para_patterns = [
    re.compile(r"\bpara(?:graph)?\.?\s*(\d{1,4})\b", re.I),
    re.compile(r"\brn\.?\s*(\d{1,4})\b", re.I),
    re.compile(r"\[\s*(\d{1,4})\s*\]"),
    re.compile(r"¶\s*(\d{1,4})"),
]

# Case markers: must have the word "Case", "Case Study"
_case_patterns = [
    re.compile(r"\bCase\s*Study\s*(\d{1,4})\b", re.I),
    re.compile(r"\bCase\s*(\d{1,4})\b", re.I),
    ]

def detect_para_numbers(text: str) -> list[int]:
    nums = []
    for pat in _para_patterns:
        matches = pat.findall(text)  # e.g., ['115'] or possibly tuples if patterns ever change
        for m in matches:
            # If a pattern produces tuples (multiple groups), pick the first non-empty
            if isinstance(m, tuple):
                m = next((g for g in m if g), "")
            # Keep only pure digits
            if not m or not m.isdigit():
                continue
            nums.append(int(m))

    # unique while preserving order
    out = []
    for n in nums:
        if n not in out:
            out.append(n)
    return out

def detect_case_numbers(text: str) -> list[int]:
    nums = []
    for pat in _case_patterns:
        nums += pat.findall(text or "")
    out = []
    for x in nums:
        n = int(x)
        if n not in out:
            out.append(n)
    return out

@st.cache_resource(show_spinner=False)

# --- Chat callbacks ------------------------------------------------------------
def clear_chat_draft():
    # Clear the persistent composer safely
    st.session_state["chat_draft"] = ""
    # optional: st.rerun()

def reset_chat():
    # Wipes the entire conversation (answers + questions + their sources)
    st.session_state["chat_history"] = []
    # optional: also clear the draft:
    # st.session_state["chat_draft"] = ""
    st.rerun()  # ensure immediate re-render

# ---------------- UI ----------------
st.set_page_config(
    page_title="EUCapML Case Tutor", 
    page_icon="⚖️", 
    layout="wide",
    initial_sidebar_state="collapsed",   # ← collapsed by default
)

# Student login
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
    
if not st.session_state.authenticated:
    with st.container():
        logo_col, title_col = st.columns([1, 5])
        with logo_col:
            st.image("assets/logo.png", width=240)
        with title_col:
            st.title("EUCapML Case Tutor")
    
    pin_input = st.text_input("Enter your password", type="password")

    try:
        correct_pin = st.secrets["STUDENT_PIN"]
    except KeyError:
        st.error("STUDENT_PIN not found in secrets. Please configure it in .streamlit/secrets.toml.")
        st.stop()

    if pin_input == correct_pin:
        st.session_state.authenticated = True
        st.success("PIN accepted. By clicking CONTINUE below you accept that this tool uses artificial intelligence and large language models, and that accordingly, answers may not be accurate. No liability is accepted for use of this tool.")
        if st.button("Continue"):
            st.rerun()
    elif pin_input:
        st.error("Incorrect PIN. Please try again.")
    st.stop()

# Sidebar (visible to all users after login)
with st.sidebar:
    st.caption(f"App version: `{APP_HASH}`")
    st.header("⚙️ Settings")
    api_key = (st.secrets.get("GROQ_API_KEY") if hasattr(st, "secrets") else None) or os.getenv("GROQ_API_KEY")
    if api_key:
        st.text_input("GROQ API Key", value="Provided via secrets/env", type="password", disabled=True)
    else:
        api_key = st.text_input("GROQ API Key", type="password", help="Set GROQ_API_KEY in Streamlit Secrets for production.")
    model_name = st.selectbox(
        "Model (free)",
        options=["llama-3.1-8b-instant", "llama-3.3-70b-versatile"],
        index=0,
        help="Both are free; 8B is faster, 70B is smarter (and slower)."
    )
    SELECTED_MODEL = model_name
    temp = st.slider("Temperature", 0.0, 1.0, 0.2, 0.05)

    st.header("🌐 Web Retrieval")
    enable_web = st.checkbox("Enable web grounding", value=True)
    max_sources = st.slider("Max sources to cite", 3, 10, 6, 1)
    st.caption("DuckDuckGo HTML + filters to EUR‑Lex, CURIA, ESMA, BaFin, Gesetze‑im‑Internet, BGH.")

    st.divider()
    st.subheader("Diagnostics")
    if st.checkbox("Run Groq connectivity test"):
        try:
            r = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key or ''}", "Content-Type": "application/json"},
                json={"model": model_name, "messages": [{"role": "user", "content": "Say: hello from Groq test"}], "max_tokens": 8},
                timeout=20,
            )
            st.write("POST /chat/completions →", r.status_code)
            st.code((r.text or "")[:1000], language="json")
        except Exception as e:
            st.exception(e)

    # --- Sidebar: Booklet Inspector (Word) ---
    st.divider()
    st.subheader("📄 Booklet Inspector")
    
    docx_source = BOOKLET  # uses your constant; you could also wire in an uploader override
    
    try:
        records, _chunks, _metas = load_booklet_anchors(docx_source)
        total = len(records)
        if total == 0:
            st.info("No anchors found in the booklet.")
        else:
            anchors_per_page = st.number_input(
                "Anchors per virtual page",
                min_value=5, max_value=50, value=12, step=1,
                help="Virtual page size = how many anchors (paras/cases) per page."
            )
            max_pages = max(1, math.ceil(total / anchors_per_page))
            virt_page = st.number_input(
                "Virtual page #",
                min_value=1, max_value=max_pages, value=1, step=1
            )
    
            start = (virt_page - 1) * anchors_per_page
            end = min(start + anchors_per_page, total)
            st.caption(f"Showing anchors {start + 1}–{end} of {total}")
    
            # Render anchors on this virtual page
            for r in records[start:end]:
                if r["kind"] == "para":
                    label = f"para {r['id']}"
                elif r["kind"] == "case":
                    label = f"Case Study {r['id']}"
                else:
                    label = "—"
                st.markdown(f"- **{label}** — {r['preview']}…")
    
            # Quick lookup for a specific para or case
            st.markdown("**Lookup**")
            lookup = st.text_input("Type e.g. `para 115` or `case 30`")
            if lookup:
                m_para = re.match(r"^\s*para\s+(\d{1,4})\s*$", lookup, re.I)
                m_case = re.match(r"^\s*case\s+(\d{1,4})\s*$", lookup, re.I)
                target_kind, target_id = None, None
                if m_para:
                    target_kind, target_id = "para", int(m_para.group(1))
                elif m_case:
                    target_kind, target_id = "case", int(m_case.group(1))
    
                if target_kind:
                    hits = [r for r in records if r["kind"] == target_kind and r["id"] == target_id]
                    if not hits:
                        st.warning(f"No match for {lookup.strip()}.")
                    else:
                        st.success(f"Found {len(hits)} match(es):")
                        for h in hits[:20]:  # cap display
                            st.markdown(f"- **#{h['idx']}** {target_kind} {h['id']} — {h['preview']}…")
                else:
                    st.info("Enter `para N` or `case K`.")
    except FileNotFoundError:
        st.error(f"Booklet not found at: {docx_source}")
    except Exception as e:
        st.exception(e)

# Main UI
st.image("assets/logo.png", width=240)
st.title("EUCapML Case Tutor")

with st.expander("📚 Case (click to read)"):
    st.write(CASE)

selected_question = st.selectbox(
    "Which question are you answering?",
    options=["Question 1", "Question 2", "Question 3"],
    index=0,
    help="This limits feedback to the selected question only."
)
st.session_state["selected_question"] = selected_question

st.subheader("📝 Your Answer")
student_answer = st.text_area("Write your solution here (≥ ~120 words).", height=260)

# ------------- Actions -------------
colA, colB = st.columns([1, 1])

with colA:
    if st.button("🔎 Generate Feedback"):
        if len(student_answer.strip()) < 80:
            st.warning("Please write a bit more so I can evaluate meaningfully (≥ 80 words).")
        else:
            with st.spinner("Scoring and collecting sources..."):
                backend = load_embedder()
                model_answer_filtered, extracted_issues = filter_model_answer_and_rubric(selected_question, MODEL_ANSWER, api_key)
                extracted_keywords = [kw for issue in extracted_issues for kw in issue.get("keywords", [])]
                rubric = generate_rubric_from_model_answer(
                    student_answer,
                    model_answer_filtered,
                    backend,
                    api_key,
                    DEFAULT_WEIGHTS
                )
                
                top_pages, source_lines = [], []
                if enable_web:
                    pages = collect_corpus(student_answer, extracted_keywords, "", max_fetch=18)
                    if not pages:
                        # fall back to “no sources” quickly
                        top_pages, source_lines = [], []
                    else:
                        top_pages, source_lines = retrieve_snippets_with_booklet(
                            student_answer, model_answer_filtered, pages, backend, extracted_keywords,
                            user_query="", top_k_pages=max_sources, chunk_words=170
                        )                    
                    
            # Breakdown
            with st.expander("🔬 Issue-by-issue breakdown"):
                for row in rubric["per_issue"]:
                    st.markdown(f"**{row['issue']}** — {row['score']} / {row['max_points']}")
                    st.markdown(f"- ✅ Found: {', '.join(row['keywords_hit']) if row['keywords_hit'] else '—'}")
                    miss = [kw for kw in row["keywords_total"] if kw not in row["keywords_hit"]]
                    st.markdown(f"- ⛔ Missing: {', '.join(miss) if miss else '—'}")

            # Deterministic corrections
            if rubric["substantive_flags"]:
                st.markdown("### ⚖️ Detected substantive flags")
                for fl in rubric["substantive_flags"]:
                    st.markdown(f"- ⚖️ {fl}")
            
            # LLM narrative feedback
            sources_block = "\n".join(source_lines) if source_lines else "(no web sources available)"
            excerpts_items = []
            for i, tp in enumerate(top_pages):
                for sn in tp["snippets"]:
                    excerpts_items.append(f"[{i+1}] {sn}")
            excerpts_block = "\n\n".join(excerpts_items[: max_sources * 3]) if excerpts_items else "(no excerpts)"
            
            # --- Narrative Feedback (fixed) ---
            st.markdown("### 🧭 Narrative Feedback")
            if api_key:
                # Trim large blocks *before* building the prompt
                sources_block = truncate_block(sources_block, 1200)
                excerpts_block = truncate_block(excerpts_block, 3200)
            
                # Hard rule included here with correct quoting (no stray backslash)
                hard_rule = (
                    "HARD RULE: Do not contradict the MODEL ANSWER. "
                    "If student reasoning conflicts, state the correct conclusion per the MODEL ANSWER and explain briefly.\n\n"
                )
            
                messages = [
                    {"role": "system", "content": system_guardrails()},
                    {"role": "user", "content": hard_rule + build_feedback_prompt(
                        student_answer, rubric, model_answer_filtered, sources_block, excerpts_block
                    )},
                ]
                            
                reply = generate_with_continuation(
                    messages, api_key, model_name=model_name, temperature=temp,
                    first_tokens=1200, continue_tokens=350
                )
                reply = enforce_model_consistency(
                    reply,
                    model_answer_filtered,
                    api_key,
                    model_name,
                )
                reply = tidy_empty_sections(reply)
                reply = prune_redundant_improvements(student_answer, reply, rubric)
                reply = lock_out_false_mistakes(reply, rubric)
                reply = lock_out_false_missing(reply, rubric)
                reply = format_feedback_and_filter_missing(reply, student_answer, model_answer_filtered, rubric)
                reply = re.sub(r"\[(?:n|N)\]", "", reply or "")
                
                used_idxs = parse_cited_indices(reply)
                display_source_lines = filter_sources_by_indices(source_lines, used_idxs) or source_lines
            
                if reply:
                    st.markdown(reply)
                else:
                    st.info("LLM unavailable. See corrections above and the issue breakdown.")
            else:
                st.info("No GROQ_API_KEY found in secrets/env. Deterministic scoring and corrections shown above.")
                display_source_lines = source_lines
            
            if source_lines:
                with st.expander("📚 Sources used"):
                    for line in display_source_lines:
                        st.markdown(f"- {line}")

with colB:
    st.markdown("### 💬 Tutor Chat: Ask me anything!")
    
    # --- state ---
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "chat_draft" not in st.session_state:
        st.session_state.chat_draft = ""

    # --- composer ---
    c1, c2, c3, c4 = st.columns([6, 1, 1, 2])
    with c1:
        st.text_area(
            "You can use this chat to ask for help with creating an answer, follow-up questions on feedback given by this app, etc.",
            key="chat_draft",
            height=90
        )
    with c2:
        send = st.button("Send", use_container_width=True, key="send_btn")
    with c3:
        st.button("Clear", use_container_width=True, key="clear_btn", on_click=clear_chat_draft)
    with c4:
        st.button("Reset chat", use_container_width=True, key="reset_chat_btn", on_click=reset_chat)
         
    # --- handle send: UPDATE STATE FIRST, DO NOT RENDER INLINE ---
    if send and st.session_state.chat_draft.strip():
        user_q = st.session_state.chat_draft
        st.session_state.chat_history.append({"role": "user", "content": user_q})

        with st.spinner("Retrieving sources and drafting a grounded reply..."):
            backend = load_embedder()
            top_pages, source_lines = [], []
            if enable_web:
                model_answer_filtered, extracted_issues = filter_model_answer_and_rubric(
                    st.session_state.get("selected_question", "Question 1"),
                    MODEL_ANSWER,
                    api_key
                )
                extracted_keywords = [kw for issue in extracted_issues for kw in issue.get("keywords", [])]

                pages = collect_corpus(student_answer, extracted_keywords, user_q, max_fetch=20)
                top_pages, source_lines = retrieve_snippets_with_booklet(
                    student_answer, model_answer_filtered, pages, backend, extracted_keywords,
                    user_query=user_q, top_k_pages=max_sources, chunk_words=170
                )
                                            
            sources_block = "\n".join(source_lines) if source_lines else "(no web sources available)"
            excerpts_items = []
            for i, tp in enumerate(top_pages):
                for sn in tp["snippets"]:
                    excerpts_items.append(f"[{i+1}] {sn}")
            excerpts_block = "\n\n".join(excerpts_items[: max_sources * 3]) if excerpts_items else "(no excerpts)"
            
            # ✅ Trim large blocks BEFORE building the prompt to free tokens for the answer
            sources_block  = truncate_block(sources_block, 1200)
            excerpts_block = truncate_block(excerpts_block, 3200)

            if api_key:
                msgs = build_chat_messages(
                    st.session_state.chat_history,
                    model_answer_filtered,
                    sources_block,
                    excerpts_block
                )
                reply = generate_with_continuation(
                    msgs, api_key, model_name=model_name, temperature=temp,
                    first_tokens=1200, continue_tokens=350
                )
                
                reply = enforce_model_consistency(
                    reply,
                    model_answer_filtered,
                    api_key,
                    model_name,    
                )

                msgs.append({"role": "system", "content":
                    "When the student's view aligns with the MODEL ANSWER, avoid marking claims as incorrect; "
                    "present extra provisions and edge cases under a short 'Suggestions' or 'Further Considerations' section."
                })
                
                # cleanup + source filtering remain unchanged
                reply = re.sub(r"\[(?:n|N)\]", "", reply or "")
                used_idxs = parse_cited_indices(reply)
                msg_sources = filter_sources_by_indices(source_lines, used_idxs) or source_lines[:]
            
            else:
                reply = None
            if not reply:
                reply = (
                    "I couldn’t reach the LLM. Here are the most relevant source snippets:\n\n"
                    + (excerpts_block if excerpts_block != "(no excerpts)" else "— no sources available —")
                    + "\n\nIn doubt, follow the model answer."
                )
        # ---- SAFETY NET (CHAT): normalize citations + show only cited sources ----
        # Safety net: remove literal “[n]”
        reply = re.sub(r"\[(?:n|N)\]", "", reply or "")
        
        # Keep only sources actually cited in this reply
        used_idxs = parse_cited_indices(reply)
        msg_sources = filter_sources_by_indices(source_lines, used_idxs) or source_lines[:]

        st.session_state.chat_history.append({
            "role": "assistant",
            "content": reply,
            "sources": msg_sources
        })
                
    # --- render FULL history AFTER updates so latest sources appear immediately ---
    for msg in st.session_state.chat_history:
        role = msg.get("role", "")
        if role in ("user", "assistant"):
            with st.chat_message(role):
                st.write(msg.get("content", ""))
                if role == "assistant":
                    # Per-message "Sources used"
                    st.markdown("#### 📚 Sources used")
                    srcs = msg.get("sources", [])
                    if not srcs:
                        st.write("— no web sources available —")
                    else:
                        for line in srcs:
                            st.markdown(f"- {line}")

st.divider()
st.markdown(
    "ℹ️ **Notes**: (c) 2025 by Stephan Balthasar. This app provides feedback based on artificial intelligence and large language models, and as a result, answers can be inaccurate. " 
    "Students are advised to use caution when using the feedback engine and chat functions. App feedback must not be read as an indicator for grades in a real examination."
)
